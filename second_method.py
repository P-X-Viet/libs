import re
from docx import Document

def analyze_ini_file(ini_file_path):
    data = {}
    current_section = None

    with open(ini_file_path, 'r') as file:
        for line in file:
            line = line.strip()

            # Match sections like [abc] and [abc:children]
            section_match = re.match(r'\[([a-zA-Z0-9]+)(:children)?\]', line)
            if section_match:
                section_name = section_match.group(1)
                is_children = section_match.group(2) is not None
                
                if section_name not in data:
                    data[section_name] = {'count': 0, 'children': {}}

                if is_children:
                    child_section = f"{section_name}:children"
                    data[section_name]['children'][child_section] = 0
                    current_section = child_section
                else:
                    current_section = section_name
            else:
                # Count each line under the current section
                if current_section:
                    if ':children' in current_section:
                        parent_section = current_section.split(':')[0]
                        data[parent_section]['children'][current_section] += 1
                    else:
                        data[current_section]['count'] += 1

    # Now we need to loop through the data to count corresponding lines for children
    for section in data.keys():
        for child in data[section]['children']:
            # Count lines for each child section
            child_count = data[section]['children'][child]
            # Loop through the file again to count lines for this child section
            with open(ini_file_path, 'r') as file:
                for line in file:
                    line = line.strip()
                    if line.startswith(f'[{section}]'):
                        # Check if it is a child line
                        child_line_match = re.match(rf'\[{section}:children\]', line)
                        if child_line_match:
                            child_count += 1
            
            # Update the count for the child
            data[section]['children'][child] = child_count

    return data

def write_to_docx(results, output_file_path):
    doc = Document()
    doc.add_heading('INI File Analysis', 0)

    for section, info in results.items():
        doc.add_paragraph(f"Section [{section}] has {info['count']} related lines.")
        for child, count in info['children'].items():
            doc.add_paragraph(f"  - Child Section [{child}] has {count} related lines.")

    doc.save(output_file_path)

def main():
    ini_file_path = 'example.ini'  # Your .ini file path
    output_file_path = 'analysis_results.docx'  # Output DOCX file path

    results = analyze_ini_file(ini_file_path)
    write_to_docx(results, output_file_path)
    print(f"Results saved to {output_file_path}")

if __name__ == '__main__':
    main()
